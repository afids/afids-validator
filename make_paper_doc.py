"""Generate a Word (.docx) version of the manuscript from the Markdown source.

Parses ``AFIDs_Validator_Aperture_Education_AT_FINAL.md`` (the curated draft) and
renders it to ``AFIDs_Validator_Aperture_Education_AT_FINAL.docx`` with headings,
inline formatting, tables, and lists preserved, then embeds the figure set from
``paper_figures/`` at the end. Because it reads the Markdown directly, the Word
output never drifts from the source.

    python make_paper_doc.py

Requires python-docx (`pip install python-docx`).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent
SRC = REPO / "AFIDs_Validator_Aperture_Education_AT_FINAL.md"
OUT = SRC.with_suffix(".docx")
FIG_DIR = REPO / "paper_figures"

# Graphical abstract + Figures 1–6, in manuscript order.
FIGURES = [
    ("Graphical abstract", "fig0_graphical_abstract.png"),
    ("Figure 1", "fig1_interface.png"),
    ("Figure 2", "fig2_learning_cycle.png"),
    ("Figure 3", "fig3_landmarks.png"),
    ("Figure 4", "fig4_qc_catch.png"),
    ("Figure 5", "fig5_difficulty.png"),
    ("Figure 6", "fig6_two_difficulties.png"),
]

BACKMATTER = {
    "Code and Data Availability",
    "Acknowledgements",
    "Author Contributions",
    "Competing Interests",
    "References",
    "Figure Legends",
}

LINK_BLUE = RGBColor(0x0B, 0x57, 0xD0)
PLACEHOLDER_RED = RGBColor(0xC0, 0x00, 0x00)

_INLINE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<ital>[^*]+?)\*"
    r"|`(?P<code>[^`]+?)`"
    r"|\[(?P<ltext>[^\]]+?)\]\((?P<lurl>[^)]+?)\)"
)
_UNESCAPE = re.compile(r"\\([\\`*_{}\[\]()>#+.!~&<-])")


def _unescape(s: str) -> str:
    return _UNESCAPE.sub(r"\1", s)


def _add_run(
    p, text, size, *, bold=False, italic=False, mono=False, link=False
):
    if not text:
        return
    r = p.add_run(_unescape(text))
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if mono:
        r.font.name = "Consolas"
    if link:
        r.font.color.rgb = LINK_BLUE
    # flag unfilled author-supplied placeholders in red so they are easy to find
    if text.strip().startswith(("Funding sources", "To be completed")):
        r.font.color.rgb = PLACEHOLDER_RED
        r.font.italic = True


def _add_inline(p, text, size=11):
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _add_run(p, text[pos : m.start()], size)
        if m.group("bold") is not None:
            _add_run(p, m.group("bold"), size, bold=True)
        elif m.group("ital") is not None:
            _add_run(p, m.group("ital"), size, italic=True)
        elif m.group("code") is not None:
            _add_run(p, m.group("code"), size, mono=True)
        elif m.group("ltext") is not None:
            _add_run(p, m.group("ltext"), size, link=True)
        pos = m.end()
    if pos < len(text):
        _add_run(p, text[pos:], size)


def _is_table_row(line: str) -> bool:
    return line.lstrip().startswith("|")


def _is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def _cells(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, block):
    rows = [_cells(ln) for ln in block if not _is_separator(ln)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline(para, row[ci] if ci < len(row) else "", size=10)
            if ri == 0:
                for run in para.runs:
                    run.font.bold = True
    doc.add_paragraph()


def build():
    if not SRC.exists():
        raise SystemExit(f"Source markdown not found: {SRC}")
    lines = SRC.read_text().splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    seen_section = False  # have we passed the front-matter (### author block)?
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # Markdown table block
        if _is_table_row(raw):
            block = []
            while i < n and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            _add_table(doc, block)
            continue

        # Title (single #)
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            txt = line[2:].strip().strip("*")
            r = p.add_run(_unescape(txt))
            r.font.bold = True
            r.font.size = Pt(16)
            i += 1
            continue

        # Section heading (## **N. Title**)
        if line.startswith("## "):
            seen_section = True
            txt = re.sub(r"^\*+|\*+$", "", line[3:].strip())
            doc.add_heading(_unescape(txt), level=1)
            i += 1
            continue

        # ### lines: author block before first section, else subheading
        if line.startswith("### "):
            txt = re.sub(r"^\*+|\*+$", "", line[4:].strip())
            if not txt:
                i += 1
                continue
            if not seen_section:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_inline(p, line[4:].strip(), size=11)
            else:
                doc.add_heading(_unescape(txt), level=2)
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:].strip())
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue

        # Back-matter section labels rendered as headings
        if line in BACKMATTER:
            doc.add_heading(line, level=1)
            i += 1
            continue

        # Ordinary paragraph
        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1

    _append_figures(doc)
    doc.save(OUT)
    print(f"Saved: {OUT}")


def _append_figures(doc):
    present = [
        (cap, FIG_DIR / fn) for cap, fn in FIGURES if (FIG_DIR / fn).exists()
    ]
    if not present:
        return
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for cap, path in present:
        c = doc.add_paragraph()
        run = c.add_run(cap)
        run.font.bold = True
        run.font.size = Pt(11)
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(path), width=Inches(6.0))
        doc.add_paragraph()
    missing = [fn for cap, fn in FIGURES if not (FIG_DIR / fn).exists()]
    if missing:
        print(
            f"  (note: {len(missing)} figure PNG(s) absent, skipped: {', '.join(missing)})"
        )


if __name__ == "__main__":
    build()
